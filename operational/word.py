import pandas as pd
from docx import Document

csv_file = '../results/results.csv'
df = pd.read_csv(csv_file)
doc = Document()
doc.add_table(rows=1, cols=len(df.columns)-1, style='Table Grid')
header_row = doc.tables[0].rows[0].cells

target_col = -1
for col_idx, column in enumerate(df.columns):
    if column == "selected_features":
        target_col = col_idx
        continue
    header_row[col_idx].text = str(column)

for _, row in df.iterrows():
    new_row = doc.tables[0].add_row().cells
    for col_idx, value in enumerate(row):
        if col_idx == target_col:
            continue
        new_row[col_idx].text = str(value)

doc.save('output_document.docx')
